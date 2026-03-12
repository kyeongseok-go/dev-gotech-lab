#!/usr/bin/env python3
"""gotech.lab 1주 MVP 스프린트 요약 Word 문서 생성"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# --- 스타일 설정 ---
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

# 제목 스타일
for i in range(1, 4):
    heading = doc.styles[f"Heading {i}"]
    heading.font.name = "맑은 고딕"
    heading.font.color.rgb = RGBColor(0x1A, 0x24, 0x3B)

doc.styles["Heading 1"].font.size = Pt(24)
doc.styles["Heading 2"].font.size = Pt(16)
doc.styles["Heading 3"].font.size = Pt(13)

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_table(doc, headers, rows, col_widths=None):
    """스타일링된 테이블 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 데이터
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_bullet_list(doc, items, bold_prefix=False):
    """불릿 리스트 추가"""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix and ":" in item:
            parts = item.split(":", 1)
            run = p.add_run(parts[0] + ":")
            run.bold = True
            run.font.size = Pt(10.5)
            run2 = p.add_run(parts[1])
            run2.font.size = Pt(10.5)
        else:
            run = p.add_run(item)
            run.font.size = Pt(10.5)


# ============================================================
# 표지
# ============================================================
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_heading("gotech.lab", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x24, 0x3B)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("1주 MVP 스프린트 결과 보고서")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x60, 0xA5, 0xFA)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("개인 홈페이지 · 기술 블로그 · 서비스 플랫폼")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("2026년 3월  |  Claude Code + Next.js + Cloudflare")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()


# ============================================================
# 1. 프로젝트 개요
# ============================================================
doc.add_heading("1. 프로젝트 개요", level=1)

doc.add_heading("1.1 목적", level=2)
doc.add_paragraph(
    "gotech.lab은 5년간의 오피스 SW 엔진 개발 경험과 AI를 활용한 빠른 프로토타이핑 과정을 "
    "기록하고 공유하는 개인 기술 플랫폼입니다. "
    "Claude Code를 AI 페어 프로그래머로 활용하여 1주 만에 MVP를 완성하는 것이 목표였습니다."
)

doc.add_heading("1.2 대상 사용자", level=2)
add_bullet_list(doc, [
    "채용 담당자: 포트폴리오 및 기술 역량 확인",
    "기술 리더: 기술 블로그 및 프로젝트 구조 검토",
    "개발자 커뮤니티: AI 실험 결과 및 기술 인사이트 공유",
    "일반 방문자: 서비스 이용 및 구독",
], bold_prefix=True)

doc.add_heading("1.3 핵심 수치", level=2)
add_table(doc, ["항목", "수치"], [
    ["스프린트 기간", "7일 (Day 0 ~ Day 7)"],
    ["페이지 라우트", "12개"],
    ["MDX 콘텐츠", "8개 (블로그 3 + 프로젝트 2 + Showcase 3)"],
    ["D1 테이블", "3개 (service_registry, news_cache, feature_flags)"],
    ["커밋 수", "Day 0~7 전체 커밋 포함"],
], col_widths=[5, 12])


# ============================================================
# 2. 기술 스택
# ============================================================
doc.add_heading("2. 기술 스택", level=1)

add_table(doc, ["영역", "기술", "버전/비고"], [
    ["프레임워크", "Next.js (App Router) + TypeScript", "v16.1.5"],
    ["UI", "Tailwind CSS + shadcn/ui", "v4 + radix-nova"],
    ["콘텐츠", "MDX + Velite", "빌드 타임 컴파일"],
    ["코드 하이라이팅", "rehype-pretty-code + Shiki", "github-dark-default"],
    ["데이터베이스", "Cloudflare D1 (SQLite)", "3개 테이블"],
    ["스토리지", "Cloudflare R2", "헬퍼 구현 완료"],
    ["배포", "Cloudflare Pages + Workers", "OpenNext adapter"],
    ["패키지 매니저", "pnpm", "9+"],
    ["AI 도구", "Claude Code (Opus 4.6)", "페어 프로그래밍"],
], col_widths=[4, 6, 6])


# ============================================================
# 3. 스프린트 타임라인
# ============================================================
doc.add_heading("3. 스프린트 타임라인", level=1)

days = [
    ("Day 0", "초기화", [
        "GitHub 저장소 생성 및 초기 커밋",
        "Claude Code 설정 + CLAUDE.md 규칙 파일",
        "docs/plan.md, docs/techspec.md 구조 수립",
        "Cloudflare 인증 초기화",
    ]),
    ("Day 1", "프로젝트 기본 구성", [
        "Next.js 16 + Cloudflare + Tailwind CSS v4 + shadcn/ui 초기화",
        "ESLint next lint → eslint CLI 마이그레이션 (Next.js 16 대응)",
        "SiteHeader / SiteFooter / PageContainer 레이아웃 컴포넌트 생성",
        "6개 페이지 라우트 스텁: /, /blog, /blog/[slug], /projects, /showcase, /about",
        "pnpm preview (Cloudflare Workers 로컬) 기동 확인",
    ]),
    ("Day 2", "콘텐츠 파이프라인", [
        "Velite 설치 및 3개 컬렉션 스키마 정의 (blog/projects/showcase)",
        "s.mdx() 기반 MDX 컴파일 + MDXContent 렌더링 컴포넌트",
        "src/lib/content.ts 조회 유틸 (getPublishedBlogs, getBlogBySlug 등)",
        "/blog 목록·상세 (SSG 3개), /projects 목록·상세 (SSG 2개), /showcase 목록 (3개)",
        "홈 최신 콘텐츠 섹션 (blog 3 / projects 2 / showcase 3 + 더보기 링크)",
    ]),
    ("Day 3", "블로그 고도화", [
        "rehype-pretty-code + Shiki 코드 하이라이팅 (github-dark-default 테마)",
        "MDX 커스텀 컴포넌트: Callout(info/warning/tip), Table 스타일링, Image(figure) 래퍼",
        "읽기 시간 계산 (getReadingTime) + 이전/다음 글 네비게이션",
        "h2/h3 기반 정적 TOC (extractToc + Toc 컴포넌트)",
        "/blog 카테고리·태그 필터 (URL searchParams 기반, BlogFilter 클라이언트 컴포넌트)",
    ]),
    ("Day 4", "페이지 확장", [
        "/about 실내용: 경력 타임라인(2019–2024), 기술 스택 4개 카테고리, 관심 분야, 사이트 소개",
        "SubscribeForm 재사용 컴포넌트 (idle/submitting/success/error 상태)",
        "/subscribe 독립 페이지 + /about 하단에도 배치",
        "/showcase/[slug] 상세 페이지 (SSG, status 뱃지, 외부 링크, 목록 네비게이션)",
        "프로젝트 상세 메타 보강 (techStack 뱃지, repo/demo 링크 버튼, 목록 네비게이션)",
    ]),
    ("Day 5", "D1 + R2 + 서비스", [
        "D1 마이그레이션: service_registry, news_cache, feature_flags 3개 테이블",
        "시드 데이터: 서비스 2개, 뉴스 3개, 피처 플래그 3개",
        "D1 접근 레이어: src/lib/db/ (services.ts, news.ts, flags.ts — try/catch fallback)",
        "R2 바인딩 + src/lib/storage/ (경로 빌더 + put/get/delete/head 헬퍼)",
        "/services 목록 페이지 (D1 기반, 상태 뱃지, 헤더 네비게이션 추가)",
        "/services/news 뉴스 애그리게이터 MVP (카테고리 필터, fallback UI)",
    ]),
    ("Day 6", "홈 페이지 폴리싱", [
        "2D 히어로 섹션: 도트 패턴 배경 + 메인 카피 + 키워드 태그 + CTA 버튼 2개",
        "키워드 태그 CSS float 애니메이션 (prefers-reduced-motion 대응)",
        "홈 섹션 정리: 프로젝트·Showcase 카드에 상세 페이지 Link 추가",
        "서비스 섹션 추가 (D1 getServices 기반, fallback UI)",
        "구독 CTA 섹션 (SubscribeForm 포함, 하단 배치)",
        "광고 placeholder 컴포넌트 (AdPlaceholder, 블로그↔프로젝트 사이)",
        "3D는 성능·일정 리스크 판단 후 의도적 제외 → Phase 2",
    ]),
    ("Day 7", "SEO + QA + 마감", [
        "전역 metadata: metadataBase, title.template (\"%s | gotech.lab\"), OG, twitter",
        "페이지별 metadata: 정적 8개 + generateMetadata 3개 (blog/projects/showcase 상세)",
        "sitemap.ts: 정적 6경로 + 동적 상세 페이지 (Velite 기반, draft 제외)",
        "robots.ts: 전체 허용, /api/ 차단, sitemap URL 명시",
        "RSS 2.0: /rss.xml Route Handler (blog 컬렉션 기반, XML escape, 1시간 캐시)",
        "README.md 전면 재작성 (기술 스택, 실행 방법, 폴더 구조, Phase 2 계획)",
        ".env.example 추가, src/lib/constants.ts 사이트 상수 분리",
        "최종 QA: 모바일 네비게이션 간격 조정, ESLint .wrangler ignore 추가",
    ]),
]

for day, title, items in days:
    doc.add_heading(f"3.{days.index((day, title, items)) + 1} {day}: {title}", level=2)
    add_bullet_list(doc, items)


# ============================================================
# 4. 사이트 구조
# ============================================================
doc.add_heading("4. 사이트 구조", level=1)

doc.add_paragraph("전체 12개 페이지 라우트로 구성되어 있습니다.")

add_table(doc, ["경로", "페이지", "렌더링", "주요 기능"], [
    ["/", "홈", "Static", "히어로 + 블로그/프로젝트/Showcase/서비스/구독 섹션"],
    ["/blog", "블로그 목록", "Dynamic", "카테고리·태그 필터 (searchParams)"],
    ["/blog/[slug]", "블로그 상세", "SSG", "MDX 본문, TOC, 코드 하이라이팅, 이전/다음"],
    ["/projects", "프로젝트 목록", "Static", "카드 목록 + 상세 링크"],
    ["/projects/[slug]", "프로젝트 상세", "SSG", "techStack 뱃지, repo/demo 링크"],
    ["/showcase", "Showcase 목록", "Static", "상태 뱃지 + 카드"],
    ["/showcase/[slug]", "Showcase 상세", "SSG", "스택, 외부 링크, MDX 본문"],
    ["/services", "서비스 목록", "Static", "D1 기반, 상태 뱃지"],
    ["/services/news", "뉴스", "Dynamic", "D1 news_cache, 카테고리 필터"],
    ["/about", "소개", "Static", "경력 타임라인, 기술 스택, 구독"],
    ["/subscribe", "구독", "Static", "SubscribeForm (더미 전송)"],
], col_widths=[3.5, 2.5, 2, 8])

doc.add_paragraph("추가 엔드포인트:")
add_bullet_list(doc, [
    "/sitemap.xml — 정적 + 동적 경로 sitemap",
    "/robots.txt — 크롤러 규칙",
    "/rss.xml — 블로그 RSS 2.0 피드",
])


# ============================================================
# 5. 데이터 아키텍처
# ============================================================
doc.add_heading("5. 데이터 아키텍처", level=1)

doc.add_paragraph("데이터는 역할에 따라 4개 저장소로 분리됩니다.")

add_table(doc, ["저장소", "역할", "현재 상태"], [
    ["Velite (MDX)", "블로그, 프로젝트, Showcase 원본 콘텐츠", "사용 중 (8개 파일)"],
    ["Cloudflare D1", "서비스 목록, 뉴스 캐시, 피처 플래그", "사용 중 (3개 테이블)"],
    ["Cloudflare R2", "이미지, 썸네일, OG 이미지", "헬퍼 구현 완료, Phase 2 실사용"],
    ["Supabase", "인증, 댓글, 구독, 멤버십", "Phase 2 예정"],
], col_widths=[4, 6, 6])

doc.add_heading("5.1 D1 테이블 구조", level=2)
add_table(doc, ["테이블", "주요 컬럼", "용도"], [
    ["service_registry", "slug, title, status, url, featured", "서비스 목록 관리"],
    ["news_cache", "title, url, category, source, published_at", "뉴스 캐시"],
    ["feature_flags", "key, value, enabled", "기능 플래그"],
], col_widths=[4, 6, 6])

doc.add_heading("5.2 R2 경로 규칙", level=2)
add_bullet_list(doc, [
    "blog/{slug}/cover.{ext} — 블로그 커버 이미지",
    "projects/{slug}/thumbnail.{ext} — 프로젝트 썸네일",
    "og/{type}/{slug}.png — OG 이미지",
])


# ============================================================
# 6. AI 페어 프로그래밍
# ============================================================
doc.add_heading("6. AI 페어 프로그래밍 워크플로우", level=1)

doc.add_paragraph(
    "Claude Code(Opus 4.6)를 AI 페어 프로그래머로 활용하여 체계적인 개발 프로세스를 적용했습니다. "
    "매일 아래 5단계를 반복하여 안정적으로 기능을 누적했습니다."
)

add_table(doc, ["단계", "활동", "산출물"], [
    ["1. 문서 읽기", "CLAUDE.md, plan.md, techspec.md 확인", "작업 컨텍스트 파악"],
    ["2. 계획 수립", "작업 범위, 수정 파일 목록, 전략 제시", "구현 계획 (코드 전 설명)"],
    ["3. 구현", "코드 생성, 파일 수정", "동작하는 코드"],
    ["4. 검증", "pnpm lint + pnpm build + 수동 체크리스트", "빌드 통과 확인"],
    ["5. 커밋", "상세 커밋 메시지, GitHub push, plan.md 갱신", "이력 관리"],
], col_widths=[3, 5.5, 5.5])

doc.add_heading("6.1 규칙 파일 체계", level=2)
add_bullet_list(doc, [
    "CLAUDE.md: 프로젝트 운영 지침 (진입점)",
    "docs/claude/project.md: 프로젝트 개요, 기술 스택, 품질 목표",
    "docs/claude/workflow.md: 작업 흐름, 구현 원칙, 검증 절차",
    "docs/claude/architecture.md: 데이터 역할 분리, 페이지 우선순위, 금지 사항",
    "docs/claude/mistakes.md: 반복 실수 추적",
    "docs/plan.md: 일별 작업 체크리스트",
])


# ============================================================
# 7. 폴더 구조
# ============================================================
doc.add_heading("7. 폴더 구조", level=1)

structure = """src/
  app/                  # App Router 페이지 및 라우트
    blog/               # 블로그 목록/상세
    projects/           # 프로젝트 목록/상세
    showcase/           # AI Showcase 목록/상세
    services/           # 서비스 목록, 뉴스
    about/              # 소개
    subscribe/          # 구독
    sitemap.ts          # 사이트맵
    robots.ts           # robots.txt
    rss.xml/            # RSS 피드
  components/
    layout/             # SiteHeader, SiteFooter, PageContainer
    blog/               # BlogFilter
    mdx/                # MDXContent, Callout, Toc
  lib/
    constants.ts        # 사이트 상수 (URL, 이름)
    content.ts          # Velite 컬렉션 조회 유틸
    db/                 # D1 접근 레이어
    storage/            # R2 헬퍼
content/                # MDX 콘텐츠 원본
migrations/             # D1 SQL 마이그레이션"""

p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = "Consolas"
run.font.size = Pt(9.5)


# ============================================================
# 8. 향후 계획
# ============================================================
doc.add_heading("8. 향후 계획 (Phase 2)", level=1)

doc.add_paragraph("1주 스프린트에서 의도적으로 제외한 항목과 확장 계획입니다.")

doc.add_heading("8.1 인증 · 사용자", level=2)
add_bullet_list(doc, [
    "Supabase Auth 연동",
    "댓글 시스템",
    "구독 백엔드 연동 (MailerLite / Buttondown / 자체 API)",
    "멤버십 기능",
])

doc.add_heading("8.2 콘텐츠 확장", level=2)
add_bullet_list(doc, [
    "뉴스 외부 수집 자동화 (크롤러 / RSS / 스케줄러)",
    "OG 이미지 동적 생성",
    "R2 이미지 업로드 파이프라인",
    "블로그 시리즈 기능",
])

doc.add_heading("8.3 UX 개선", level=2)
add_bullet_list(doc, [
    "다크모드 토글",
    "3D 히어로 (성능 검증 후 도입)",
    "모바일 햄버거 메뉴",
    "페이지 전환 애니메이션",
])

doc.add_heading("8.4 운영", level=2)
add_bullet_list(doc, [
    "Cloudflare 프로덕션 배포 + 커스텀 도메인",
    "Lighthouse 성능 최적화",
    "GitHub Actions CI/CD",
    "모니터링 및 에러 트래킹",
])


# ============================================================
# 저장
# ============================================================
output_path = "/Users/eunsuk.ko128/workspace/dev-gotech-lab/docs/output/gotech-lab-sprint-report.docx"
doc.save(output_path)
print(f"Word 저장 완료: {output_path}")
